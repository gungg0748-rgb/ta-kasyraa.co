"""
Generate Bab 4 (Hasil dan Pembahasan) DOCX untuk Kasyraa.co
Format: Times New Roman, spasi 2.0, indent 1cm, justify, margin 2cm
"""
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import glob

BASE_DIR = r"C:\Users\imade\OneDrive\Documents\Ngodingteruzz\Kasyraa.co"
SCREENSHOT_DIR = os.path.join(BASE_DIR, "tools", "screenshots")
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "Bab_4_Hasil_dan_Pembahasan_v3.docx")

os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

doc = Document()

# --- Page setup ---
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(0)
style.paragraph_format.space_before = Pt(0)
style.paragraph_format.line_spacing = 2.0
style.paragraph_format.first_line_indent = Cm(1)
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# Set font for East Asian too
rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rFonts.set(qn('w:cs'), 'Times New Roman')
rPr.insert(0, rFonts)

def add_heading_styled(doc, text, level=1):
    """Add heading with Times New Roman, correct size"""
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        rPr = run._element.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Times New Roman')
        rFonts.set(qn('w:hAnsi'), 'Times New Roman')
        rFonts.set(qn('w:cs'), 'Times New Roman')
        rPr.insert(0, rFonts)
        if level == 1:
            run.font.size = Pt(14)
        elif level == 2:
            run.font.size = Pt(12)
    h.paragraph_format.line_spacing = 2.0
    h.paragraph_format.first_line_indent = Cm(0)  # no indent for headings
    h.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return h

def add_para(doc, text, bold=False, indent=True, align=None):
    """Add paragraph with proper formatting"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    p.paragraph_format.line_spacing = 2.0
    if indent:
        p.paragraph_format.first_line_indent = Cm(1)
    else:
        p.paragraph_format.first_line_indent = Cm(0)
    if align:
        p.paragraph_format.alignment = align
    else:
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_image_centered(doc, image_path, width_cm=14):
    """Add image centered with caption"""
    if not os.path.exists(image_path):
        add_para(doc, f"[Screenshot tidak tersedia: {os.path.basename(image_path)}]", indent=False)
        return
    
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run()
    run.add_picture(image_path, width=Cm(width_cm))
    return p

def add_figure_caption(doc, caption_text):
    """Add figure caption (Gambar X.X ...) centered"""
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(12)
    p.paragraph_format.line_spacing = 1.0
    run = p.add_run(caption_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True
    return p

def add_table_bordered(doc, headers, rows, col_widths=None):
    """Add bordered table with header row"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header row
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.15
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
        run.bold = True
    # Data rows
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.15
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if c == 0:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(9)
    doc.add_paragraph()  # spacer
    return table

add_heading_styled(doc, 'BAB 4', level=1)
add_heading_styled(doc, 'HASIL DAN PEMBAHASAN', level=1)

add_para(doc, 'Bab ini membahas hasil implementasi sistem informasi inventori dan point of sale Kasyraa.co yang telah dikembangkan menggunakan framework Laravel 12. Pembahasan mencakup tampilan antarmuka, fungsi-fungsi utama, serta alur kerja dari setiap modul yang tersedia dalam sistem.')

# --- 4.1 Login ---
add_heading_styled(doc, '4.1  Halaman Login', level=2)
add_para(doc, 'Halaman login merupakan halaman pertama yang diakses oleh pengguna sebelum memasuki sistem. Pengguna diharuskan memasukkan alamat email dan kata sandi yang telah terdaftar. Sistem akan memverifikasi kredensial dan mengarahkan pengguna ke dashboard sesuai dengan peran (role) masing-masing: admin, kasir, atau gudang.')
add_para(doc, 'Tampilan halaman login disajikan pada Gambar 4.1.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '01-login.png'))
add_figure_caption(doc, 'Gambar 4.1 Tampilan Halaman Login')

# --- 4.2 Dashboard ---
add_heading_styled(doc, '4.2  Halaman Dashboard', level=2)
add_para(doc, 'Dashboard merupakan halaman utama setelah pengguna berhasil melakukan login. Halaman ini menampilkan ringkasan informasi penting seperti total produk, jumlah produk yang perlu restock, penjualan hari ini, pembelian hari ini, serta daftar aktivitas terbaru. Dashboard membantu pengguna memonitor kondisi bisnis secara cepat.')
add_para(doc, 'Tampilan halaman dashboard disajikan pada Gambar 4.2.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '02-dashboard.png'))
add_figure_caption(doc, 'Gambar 4.2 Tampilan Halaman Dashboard')

# --- 4.3 Kategori ---
add_heading_styled(doc, '4.3  Modul Kategori', level=2)

add_heading_styled(doc, '4.3.1  Daftar Kategori', level=3)
add_para(doc, 'Halaman daftar kategori menampilkan seluruh kategori produk yang terdaftar dalam sistem dalam bentuk tabel. Setiap baris menampilkan nama kategori, deskripsi, jumlah produk terkait, serta tombol aksi untuk mengedit atau menghapus kategori. Halaman ini hanya dapat diakses oleh admin.')
add_para(doc, 'Tampilan halaman daftar kategori disajikan pada Gambar 4.3.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '03-kategori-index.png'))
add_figure_caption(doc, 'Gambar 4.3 Tampilan Daftar Kategori')

add_heading_styled(doc, '4.3.2  Form Tambah Kategori', level=3)
add_para(doc, 'Form tambah kategori digunakan untuk menambahkan kategori produk baru ke dalam sistem. Admin mengisi nama kategori dan deskripsi, kemudian menyimpan data. Kategori yang telah ditambahkan akan muncul pada daftar kategori dan dapat digunakan untuk mengelompokkan produk.')
add_para(doc, 'Tampilan form tambah kategori disajikan pada Gambar 4.4.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '04-kategori-create.png'))
add_figure_caption(doc, 'Gambar 4.4 Tampilan Form Tambah Kategori')

add_heading_styled(doc, '4.3.3  Form Edit Kategori', level=3)
add_para(doc, 'Form edit kategori digunakan untuk mengubah data kategori yang sudah ada. Admin dapat memperbarui nama kategori dan deskripsi. Tampilan form edit serupa dengan form tambah, namun seluruh field sudah terisi dengan data kategori yang akan diubah.')
add_para(doc, 'Tampilan form edit kategori disajikan pada Gambar 4.5.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '05-kategori-edit.png'))
add_figure_caption(doc, 'Gambar 4.5 Tampilan Form Edit Kategori')

# --- 4.4 Satuan ---
add_heading_styled(doc, '4.4  Modul Satuan', level=2)

add_heading_styled(doc, '4.4.1  Daftar Satuan', level=3)
add_para(doc, 'Halaman daftar satuan menampilkan seluruh satuan barang yang digunakan dalam sistem, seperti pcs, lusin, kodi, dan sebagainya. Setiap satuan memiliki nama dan jumlah produk yang menggunakannya. Halaman ini hanya dapat diakses oleh admin.')
add_para(doc, 'Tampilan halaman daftar satuan disajikan pada Gambar 4.6.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '06-satuan-index.png'))
add_figure_caption(doc, 'Gambar 4.6 Tampilan Daftar Satuan')

add_heading_styled(doc, '4.4.2  Form Tambah Satuan', level=3)
add_para(doc, 'Form tambah satuan digunakan untuk menambahkan satuan barang baru. Admin mengisi nama satuan, kemudian menyimpan data. Satuan yang telah ditambahkan dapat langsung digunakan pada form produk.')
add_para(doc, 'Tampilan form tambah satuan disajikan pada Gambar 4.7.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '07-satuan-create.png'))
add_figure_caption(doc, 'Gambar 4.7 Tampilan Form Tambah Satuan')

add_heading_styled(doc, '4.4.3  Form Edit Satuan', level=3)
add_para(doc, 'Form edit satuan digunakan untuk mengubah nama satuan yang sudah ada. Tampilan form edit serupa dengan form tambah, namun field sudah terisi dengan data satuan yang akan diubah.')
add_para(doc, 'Tampilan form edit satuan disajikan pada Gambar 4.8.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '08-satuan-edit.png'))
add_figure_caption(doc, 'Gambar 4.8 Tampilan Form Edit Satuan')

# --- 4.5 Supplier ---
add_heading_styled(doc, '4.5  Modul Supplier', level=2)

add_heading_styled(doc, '4.5.1  Daftar Supplier', level=3)
add_para(doc, 'Halaman daftar supplier menampilkan seluruh pemasok yang terdaftar dalam sistem. Informasi yang ditampilkan meliputi nama supplier, kontak, alamat, dan jumlah pembelian yang pernah dilakukan. Halaman ini hanya dapat diakses oleh admin.')
add_para(doc, 'Tampilan halaman daftar supplier disajikan pada Gambar 4.9.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '09-supplier-index.png'))
add_figure_caption(doc, 'Gambar 4.9 Tampilan Daftar Supplier')

add_heading_styled(doc, '4.5.2  Form Tambah Supplier', level=3)
add_para(doc, 'Form tambah supplier digunakan untuk mendaftarkan pemasok baru. Admin mengisi data seperti nama, nomor telepon, alamat, dan informasi kontak lainnya. Supplier yang telah didaftarkan dapat dipilih pada saat mencatat transaksi pembelian.')
add_para(doc, 'Tampilan form tambah supplier disajikan pada Gambar 4.10.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '10-supplier-create.png'))
add_figure_caption(doc, 'Gambar 4.10 Tampilan Form Tambah Supplier')

add_heading_styled(doc, '4.5.3  Form Edit Supplier', level=3)
add_para(doc, 'Form edit supplier digunakan untuk memperbarui data pemasok yang sudah terdaftar. Admin dapat mengubah nama, email, telepon, dan alamat supplier. Tampilan form edit serupa dengan form tambah dengan data yang sudah terisi.')
add_para(doc, 'Tampilan form edit supplier disajikan pada Gambar 4.11.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '11-supplier-edit.png'))
add_figure_caption(doc, 'Gambar 4.11 Tampilan Form Edit Supplier')

add_heading_styled(doc, '4.5.4  Detail Supplier', level=3)
add_para(doc, 'Halaman detail supplier menampilkan informasi lengkap tentang seorang pemasok beserta riwayat pembelian yang pernah dilakukan dari supplier tersebut. Halaman ini membantu admin dalam melacak hubungan bisnis dengan masing-masing pemasok.')
add_para(doc, 'Tampilan halaman detail supplier disajikan pada Gambar 4.12.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '12-supplier-detail.png'))
add_figure_caption(doc, 'Gambar 4.12 Tampilan Detail Supplier')

# --- 4.6 Produk ---
add_heading_styled(doc, '4.6  Modul Produk', level=2)

add_heading_styled(doc, '4.6.1  Daftar Produk', level=3)
add_para(doc, 'Halaman daftar produk menampilkan seluruh produk yang tersedia dalam sistem. Informasi yang ditampilkan meliputi nama produk, kategori, satuan, harga jual, barcode, jumlah varian, dan total stok. Halaman ini dapat diakses oleh admin dan gudang. Terdapat fitur pencarian dan filter untuk mempermudah navigasi.')
add_para(doc, 'Tampilan halaman daftar produk disajikan pada Gambar 4.13.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '13-produk-index.png'))
add_figure_caption(doc, 'Gambar 4.13 Tampilan Daftar Produk')

add_heading_styled(doc, '4.6.2  Form Tambah Produk', level=3)
add_para(doc, 'Form tambah produk digunakan untuk menambahkan produk baru ke dalam katalog. Pengguna mengisi data seperti nama produk, kategori, satuan, harga jual, level restock, deskripsi, dan foto produk. Barcode akan dibuat otomatis oleh sistem setelah produk disimpan. Varian produk (model, warna, ukuran) dapat ditambahkan melalui halaman detail produk.')
add_para(doc, 'Tampilan form tambah produk disajikan pada Gambar 4.14.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '14-produk-create.png'))
add_figure_caption(doc, 'Gambar 4.14 Tampilan Form Tambah Produk')

add_heading_styled(doc, '4.6.3  Form Edit Produk', level=3)
add_para(doc, 'Form edit produk digunakan untuk mengubah data produk yang sudah ada. Pengguna dapat memperbarui nama, kategori, satuan, harga jual, level restock, deskripsi, dan foto produk. Tampilan form edit serupa dengan form tambah dengan data yang sudah terisi.')
add_para(doc, 'Tampilan form edit produk disajikan pada Gambar 4.15.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '15-produk-edit.png'))
add_figure_caption(doc, 'Gambar 4.15 Tampilan Form Edit Produk')

add_heading_styled(doc, '4.6.4  Detail Produk', level=3)
add_para(doc, 'Halaman detail produk menampilkan informasi lengkap tentang suatu produk termasuk varian-variannya beserta stok masing-masing varian. Pada halaman ini pengguna juga dapat menambahkan varian baru, mengedit varian yang sudah ada, atau menghapus varian.')
add_para(doc, 'Tampilan halaman detail produk disajikan pada Gambar 4.16.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '16-produk-detail.png'))
add_figure_caption(doc, 'Gambar 4.16 Tampilan Detail Produk')

# --- 4.7 Pembelian ---
add_heading_styled(doc, '4.7  Modul Pembelian', level=2)

add_heading_styled(doc, '4.7.1  Daftar Pembelian', level=3)
add_para(doc, 'Halaman daftar pembelian menampilkan seluruh transaksi pembelian yang telah dilakukan. Informasi yang ditampilkan meliputi nomor invoice, tanggal, supplier, petugas pencatat, dan total pembelian. Halaman ini dapat diakses oleh admin dan gudang. Tersedia filter berdasarkan supplier dan rentang tanggal.')
add_para(doc, 'Tampilan halaman daftar pembelian disajikan pada Gambar 4.17.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '17-pembelian-index.png'))
add_figure_caption(doc, 'Gambar 4.17 Tampilan Daftar Pembelian')

add_heading_styled(doc, '4.7.2  Form Tambah Pembelian', level=3)
add_para(doc, 'Form tambah pembelian digunakan untuk mencatat transaksi pembelian baru. Pengguna memilih supplier, menambahkan produk yang dibeli beserta jumlah dan harga, kemudian sistem menghitung total secara otomatis. Nomor invoice dibuat otomatis oleh sistem dengan format PO-YYYYMMDD-XXXXXX.')
add_para(doc, 'Tampilan form tambah pembelian disajikan pada Gambar 4.18.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '18-pembelian-create.png'))
add_figure_caption(doc, 'Gambar 4.18 Tampilan Form Tambah Pembelian')

add_heading_styled(doc, '4.7.3  Detail Pembelian', level=3)
add_para(doc, 'Halaman detail pembelian menampilkan informasi lengkap transaksi pembelian termasuk daftar item yang dibeli, harga satuan, jumlah, subtotal, dan total keseluruhan. Halaman ini juga menampilkan informasi supplier, tanggal transaksi, dan petugas pencatat.')
add_para(doc, 'Tampilan halaman detail pembelian disajikan pada Gambar 4.19.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '19-pembelian-detail.png'))
add_figure_caption(doc, 'Gambar 4.19 Tampilan Detail Pembelian')

# --- 4.8 Penjualan ---
add_heading_styled(doc, '4.8  Modul Penjualan', level=2)

add_heading_styled(doc, '4.8.1  Daftar Penjualan', level=3)
add_para(doc, 'Halaman daftar penjualan menampilkan seluruh transaksi penjualan yang telah dilakukan. Informasi yang ditampilkan meliputi nomor invoice, tanggal, kasir, dan total penjualan. Halaman ini dapat diakses oleh admin dan kasir. Tersedia filter berdasarkan rentang tanggal.')
add_para(doc, 'Tampilan halaman daftar penjualan disajikan pada Gambar 4.20.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '20-penjualan-index.png'))
add_figure_caption(doc, 'Gambar 4.20 Tampilan Daftar Penjualan')

add_heading_styled(doc, '4.8.2  Form Tambah Penjualan', level=3)
add_para(doc, 'Form tambah penjualan digunakan untuk mencatat transaksi penjualan baru. Kasir dapat mencari produk melalui daftar, pemindai barcode, atau menghubungkan perangkat mobile sebagai scanner. Setiap item ditambahkan ke keranjang, dan sistem menghitung total secara otomatis. Nomor invoice dibuat otomatis dengan format INV-YYYYMMDD-XXXXXX.')
add_para(doc, 'Tampilan form tambah penjualan disajikan pada Gambar 4.21.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '21-penjualan-create.png'))
add_figure_caption(doc, 'Gambar 4.21 Tampilan Form Tambah Penjualan')

add_heading_styled(doc, '4.8.3  Detail Penjualan', level=3)
add_para(doc, 'Halaman detail penjualan menampilkan informasi lengkap transaksi penjualan termasuk daftar item yang dijual, harga satuan, jumlah, subtotal, dan total keseluruhan. Halaman ini juga menampilkan informasi kasir dan tanggal transaksi.')
add_para(doc, 'Tampilan halaman detail penjualan disajikan pada Gambar 4.22.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '22-penjualan-detail.png'))
add_figure_caption(doc, 'Gambar 4.22 Tampilan Detail Penjualan')

# --- 4.9 Retur ---
add_heading_styled(doc, '4.9  Modul Retur', level=2)

add_heading_styled(doc, '4.9.1  Daftar Retur', level=3)
add_para(doc, 'Halaman daftar retur menampilkan seluruh transaksi retur pembelian yang telah dilakukan. Retur dilakukan ketika barang yang diterima dari supplier mengalami kerusakan atau tidak sesuai pesanan. Informasi yang ditampilkan meliputi nomor retur, tanggal, supplier terkait, referensi pembelian, dan petugas pencatat.')
add_para(doc, 'Tampilan halaman daftar retur disajikan pada Gambar 4.23.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '23-retur-index.png'))
add_figure_caption(doc, 'Gambar 4.23 Tampilan Daftar Retur')

add_heading_styled(doc, '4.9.2  Form Tambah Retur', level=3)
add_para(doc, 'Form tambah retur digunakan untuk mencatat pengembalian barang kepada supplier. Pengguna memilih transaksi pembelian terkait, kemudian memilih item yang diretur beserta jumlah dan alasan retur. Stok produk akan otomatis berkurang sesuai jumlah retur yang dicatat.')
add_para(doc, 'Tampilan form tambah retur disajikan pada Gambar 4.24.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '24-retur-create.png'))
add_figure_caption(doc, 'Gambar 4.24 Tampilan Form Tambah Retur')

add_heading_styled(doc, '4.9.3  Detail Retur', level=3)
add_para(doc, 'Halaman detail retur menampilkan informasi lengkap transaksi retur termasuk daftar item yang dikembalikan, jumlah, alasan retur, dan total item yang diretur. Halaman ini juga menampilkan informasi supplier dan transaksi pembelian terkait.')
add_para(doc, 'Tampilan halaman detail retur disajikan pada Gambar 4.25.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '25-retur-detail.png'))
add_figure_caption(doc, 'Gambar 4.25 Tampilan Detail Retur')

# --- 4.10 Stok Opname ---
add_heading_styled(doc, '4.10  Modul Stok Opname', level=2)

add_heading_styled(doc, '4.10.1  Daftar Stok Opname', level=3)
add_para(doc, 'Halaman daftar stok opname menampilkan seluruh kegiatan pengecekan stok fisik yang telah dilakukan. Stok opname dilakukan untuk memastikan kesesuaian antara stok sistem dengan stok fisik di gudang. Informasi yang ditampilkan meliputi tanggal, petugas pencatat, catatan, dan status (draft/confirmed).')
add_para(doc, 'Tampilan halaman daftar stok opname disajikan pada Gambar 4.26.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '26-opname-index.png'))
add_figure_caption(doc, 'Gambar 4.26 Tampilan Daftar Stok Opname')

add_heading_styled(doc, '4.10.2  Form Tambah Stok Opname', level=3)
add_para(doc, 'Form tambah stok opname digunakan untuk memulai kegiatan pengecekan stok. Sistem akan menampilkan seluruh produk beserta stok sistem. Pengguna memasukkan stok fisik hasil perhitungan, dan sistem menghitung selisih secara otomatis. Data tersimpan otomatis setiap perubahan dan dapat dilanjutkan sebagai draft sebelum dikonfirmasi.')
add_para(doc, 'Tampilan form tambah stok opname disajikan pada Gambar 4.27.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '27-opname-create.png'))
add_figure_caption(doc, 'Gambar 4.27 Tampilan Form Tambah Stok Opname')

add_heading_styled(doc, '4.10.3  Detail Stok Opname', level=3)
add_para(doc, 'Halaman detail stok opname menampilkan hasil pengecekan stok secara lengkap. Setiap item menampilkan stok sistem, stok fisik, selisih, dan keterangan. Setelah dikonfirmasi, stok sistem akan diperbarui sesuai stok fisik. Ringkasan selisih ditampilkan dalam bentuk kartu informasi.')
add_para(doc, 'Tampilan halaman detail stok opname disajikan pada Gambar 4.28.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '28-opname-detail.png'))
add_figure_caption(doc, 'Gambar 4.28 Tampilan Detail Stok Opname')

# --- 4.11 Laporan ---
add_heading_styled(doc, '4.11  Modul Laporan', level=2)

add_heading_styled(doc, '4.11.1  Indeks Laporan', level=3)
add_para(doc, 'Halaman indeks laporan merupakan pusat akses ke seluruh laporan yang tersedia dalam sistem. Pengguna dapat memilih jenis laporan yang ingin dilihat: laporan stok, laporan penjualan, laporan pembelian, laporan retur, atau laporan stok opname. Setiap laporan dapat difilter berdasarkan rentang tanggal dan diekspor ke format Excel.')
add_para(doc, 'Tampilan halaman indeks laporan disajikan pada Gambar 4.29.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '29-laporan-index.png'))
add_figure_caption(doc, 'Gambar 4.29 Tampilan Indeks Laporan')

# Sub laporan
laporan_subs = [
    ('30-laporan-stock', '4.11.2', 'Laporan Stok', '4.30', 'menampilkan data stok seluruh produk termasuk stok saat ini, stok minimal, dan status ketersediaan. Laporan ini membantu manajemen dalam memantau inventori dan menentukan produk yang perlu di-restock'),
    ('31-laporan-sales', '4.11.3', 'Laporan Penjualan', '4.31', 'menampilkan riwayat transaksi penjualan dalam periode tertentu. Laporan ini mencakup informasi tanggal, nomor invoice, kasir, item terjual, dan total penjualan'),
    ('32-laporan-purchases', '4.11.4', 'Laporan Pembelian', '4.32', 'menampilkan riwayat transaksi pembelian dalam periode tertentu. Laporan ini mencakup informasi tanggal, nomor invoice, supplier, item dibeli, dan total pembelian'),
    ('33-laporan-returns', '4.11.5', 'Laporan Retur', '4.33', 'menampilkan riwayat transaksi retur pembelian dalam periode tertentu. Laporan ini mencakup informasi tanggal, nomor retur, supplier terkait, item diretur, dan total item'),
    ('34-laporan-opnames', '4.11.6', 'Laporan Stok Opname', '4.34', 'menampilkan riwayat kegiatan stok opname beserta hasilnya. Laporan ini mencakup informasi tanggal, petugas, jumlah varian dicek, selisih stok, dan status'),
]

for fname, section, title, fig, desc in laporan_subs:
    add_heading_styled(doc, f'{section}  {title}', level=3)
    add_para(doc, f'Halaman {title.lower()} {desc}.')
    add_para(doc, f'Tampilan halaman {title.lower()} disajikan pada Gambar {fig}.')
    add_image_centered(doc, os.path.join(SCREENSHOT_DIR, f'{fname}.png'))
    add_figure_caption(doc, f'Gambar {fig} Tampilan {title}')

# --- 4.12 Akun ---
add_heading_styled(doc, '4.12  Modul Manajemen Akun', level=2)

add_heading_styled(doc, '4.12.1  Daftar Akun', level=3)
add_para(doc, 'Halaman daftar akun menampilkan seluruh pengguna yang terdaftar dalam sistem. Informasi yang ditampilkan meliputi nama, email, peran (admin/kasir/gudang), dan status aktif. Admin dapat mengelola akun pengguna melalui halaman ini termasuk mengaktifkan atau menonaktifkan akun.')
add_para(doc, 'Tampilan halaman daftar akun disajikan pada Gambar 4.35.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '35-akun-index.png'))
add_figure_caption(doc, 'Gambar 4.35 Tampilan Daftar Akun')

add_heading_styled(doc, '4.12.2  Form Tambah Akun', level=3)
add_para(doc, 'Form tambah akun digunakan untuk mendaftarkan pengguna baru ke dalam sistem. Admin mengisi nama lengkap, email, kata sandi, konfirmasi kata sandi, dan memilih peran pengguna (admin, kasir, atau gudang). Akun yang baru dibuat akan langsung aktif dan dapat digunakan untuk login.')
add_para(doc, 'Tampilan form tambah akun disajikan pada Gambar 4.36.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '36-akun-create.png'))
add_figure_caption(doc, 'Gambar 4.36 Tampilan Form Tambah Akun')

add_heading_styled(doc, '4.12.3  Form Edit Akun', level=3)
add_para(doc, 'Form edit akun digunakan untuk mengubah data pengguna yang sudah terdaftar. Admin dapat memperbarui nama lengkap, email, dan peran. Kata sandi bersifat opsional — hanya diisi jika ingin mengganti kata sandi pengguna. Tampilan form edit serupa dengan form tambah dengan data yang sudah terisi.')
add_para(doc, 'Tampilan form edit akun disajikan pada Gambar 4.37.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '37-akun-edit.png'))
add_figure_caption(doc, 'Gambar 4.37 Tampilan Form Edit Akun')

# --- 4.13 Profil ---
add_heading_styled(doc, '4.13  Halaman Profil', level=2)
add_para(doc, 'Halaman profil menampilkan informasi akun pengguna yang sedang login. Pengguna dapat memperbarui informasi profil seperti nama dan email, serta mengubah kata sandi. Halaman ini dapat diakses oleh semua peran pengguna.')
add_para(doc, 'Tampilan halaman profil disajikan pada Gambar 4.38.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '38-profil.png'))
add_figure_caption(doc, 'Gambar 4.38 Tampilan Halaman Profil')

# --- 4.14 Scanner Mobile ---
add_heading_styled(doc, '4.14  Scanner Mobile', level=2)
add_para(doc, 'Scanner mobile merupakan fitur yang memungkinkan perangkat seluler (smartphone) digunakan sebagai pemindai barcode untuk aplikasi desktop Kasyraa.co. Pengguna desktop membuat token koneksi, kemudian pengguna mobile mengakses URL scanner dengan token tersebut. Setelah terhubung, setiap pemindaian barcode di perangkat mobile akan langsung dikirim ke halaman penjualan di desktop.')
add_para(doc, 'Mekanisme scanner mobile bekerja dengan cara: (1) desktop menghasilkan token unik melalui endpoint /scanner/token, (2) perangkat mobile membuka halaman /scanner/{token} yang berisi antarmuka kamera, (3) setiap kali barcode dipindai, data dikirim ke server melalui endpoint /scanner/{token}/push, (4) desktop melakukan polling ke endpoint /scanner/poll/{token} untuk menerima data barcode secara real-time.')
add_para(doc, 'Tampilan halaman scanner mobile disajikan pada Gambar 4.39.')
add_image_centered(doc, os.path.join(SCREENSHOT_DIR, '39-scanner-mobile.png'))
add_figure_caption(doc, 'Gambar 4.39 Tampilan Scanner Mobile')

# ============================================================
# 4.15 BLACK BOX TESTING
# ============================================================
# --- 4.15 Black Box Testing (Single Table) ---
add_heading_styled(doc, '4.15  Pengujian Black Box', level=2)
add_para(doc, 'Pengujian black box dilakukan untuk memverifikasi fungsionalitas sistem tanpa melihat struktur kode internal. Pengujian berfokus pada input yang diberikan dan output yang dihasilkan oleh setiap modul. Setiap skenario diuji dengan memberikan data masukan tertentu dan membandingkan hasil aktual dengan hasil yang diharapkan. Hasil pengujian black box secara keseluruhan disajikan pada Tabel 4.1.')

add_table_bordered(doc,
    ['No', 'Modul', 'Skenario Pengujian', 'Hasil yang Diharapkan', 'Hasil Pengujian', 'Status'],
    [
        # === LOGIN ===
        ['1', 'Login', 'Login dengan email dan password valid', 'Berhasil login dan diarahkan ke dashboard', 'Berhasil login dan menampilkan dashboard', 'Valid'],
        ['2', 'Login', 'Login dengan email valid, password salah', 'Gagal login, menampilkan pesan kesalahan', 'Menampilkan pesan "Email atau password salah"', 'Valid'],
        ['3', 'Login', 'Login dengan email tidak terdaftar', 'Gagal login, menampilkan pesan kesalahan', 'Menampilkan pesan "Email atau password salah"', 'Valid'],
        ['4', 'Login', 'Login dengan form kosong', 'Menampilkan validasi field required', 'Menampilkan pesan validasi pada field kosong', 'Valid'],
        # === DASHBOARD ===
        ['5', 'Dashboard', 'Mengakses dashboard setelah login', 'Menampilkan ringkasan: total produk, perlu restock, penjualan & pembelian hari ini', 'Semua ringkasan tampil dengan data yang sesuai', 'Valid'],
        ['6', 'Dashboard', 'Memverifikasi data total produk sesuai database', 'Jumlah total produk = jumlah di tabel products', 'Data sesuai (22 produk, 11 jenis)', 'Valid'],
        ['7', 'Dashboard', 'Memverifikasi daftar produk hampir habis', 'Hanya menampilkan produk dengan stok ≤ reorder_level', 'Menampilkan 4 produk dengan stok 0', 'Valid'],
        ['8', 'Dashboard', 'Memverifikasi daftar aktivitas terbaru', 'Menampilkan 5 penjualan & 3 pembelian terbaru', 'Data sesuai urutan tanggal', 'Valid'],
        # === KATEGORI ===
        ['9', 'Kategori', 'Menampilkan daftar kategori', 'Semua kategori tampil dalam tabel', 'Tabel menampilkan 4 kategori', 'Valid'],
        ['10', 'Kategori', 'Menambah kategori baru', 'Kategori tersimpan dan muncul di daftar', 'Kategori baru muncul dengan 0 produk', 'Valid'],
        ['11', 'Kategori', 'Mengedit nama kategori', 'Nama kategori berubah dan tersimpan', 'Nama berhasil diubah', 'Valid'],
        ['12', 'Kategori', 'Menghapus kategori tanpa produk', 'Kategori terhapus dari daftar', 'Kategori hilang dari daftar', 'Valid'],
        ['13', 'Kategori', 'Menghapus kategori yang memiliki produk', 'Gagal atau menampilkan peringatan constraint', 'Menampilkan error terkait constraint', 'Valid'],
        # === SATUAN ===
        ['14', 'Satuan', 'Menampilkan daftar satuan', 'Semua satuan tampil dengan jumlah produk', 'Menampilkan daftar satuan lengkap', 'Valid'],
        ['15', 'Satuan', 'Menambah satuan baru', 'Satuan tersimpan dan muncul di daftar', 'Satuan baru muncul di daftar', 'Valid'],
        ['16', 'Satuan', 'Mengedit nama satuan', 'Nama satuan berubah dan tersimpan', 'Nama satuan berhasil diubah', 'Valid'],
        ['17', 'Satuan', 'Menghapus satuan tidak terpakai', 'Satuan terhapus dari daftar', 'Satuan berhasil dihapus', 'Valid'],
        # === SUPPLIER ===
        ['18', 'Supplier', 'Menampilkan daftar supplier', 'Semua supplier tampil dengan info lengkap', 'Menampilkan daftar supplier', 'Valid'],
        ['19', 'Supplier', 'Menambah supplier dengan data lengkap', 'Supplier tersimpan dan muncul di daftar', 'Supplier baru muncul di daftar', 'Valid'],
        ['20', 'Supplier', 'Mengedit data supplier', 'Data supplier berubah dan tersimpan', 'Data supplier berhasil diubah', 'Valid'],
        ['21', 'Supplier', 'Melihat detail supplier', 'Menampilkan info kontak dan riwayat pembelian', 'Halaman detail menampilkan data lengkap', 'Valid'],
        ['22', 'Supplier', 'Menghapus supplier', 'Supplier terhapus dari daftar', 'Supplier berhasil dihapus', 'Valid'],
        # === PRODUK ===
        ['23', 'Produk', 'Menampilkan daftar produk', 'Semua produk tampil dengan data lengkap', 'Menampilkan 11 produk', 'Valid'],
        ['24', 'Produk', 'Filter produk berdasarkan kategori', 'Hanya produk kategori terpilih yang tampil', 'Filter berfungsi sesuai kategori', 'Valid'],
        ['25', 'Produk', 'Pencarian produk berdasarkan nama', 'Produk dengan nama mengandung kata kunci tampil', 'Pencarian berfungsi', 'Valid'],
        ['26', 'Produk', 'Menambah produk baru dengan data lengkap', 'Produk tersimpan, barcode auto-generated', 'Produk baru tampil dengan barcode otomatis', 'Valid'],
        ['27', 'Produk', 'Mengedit data produk', 'Data produk berubah dan tersimpan', 'Data produk berhasil diubah', 'Valid'],
        ['28', 'Produk', 'Melihat detail produk dan varian', 'Menampilkan info produk + tabel varian', 'Detail produk menampilkan data lengkap', 'Valid'],
        ['29', 'Produk', 'Menambah varian baru pada produk', 'Varian tersimpan dan muncul di tabel', 'Varian baru ditambahkan', 'Valid'],
        ['30', 'Produk', 'Mengedit varian produk', 'Data varian berubah dan tersimpan', 'Data varian berhasil diubah', 'Valid'],
        ['31', 'Produk', 'Menghapus varian produk', 'Varian terhapus dari tabel', 'Varian berhasil dihapus', 'Valid'],
        ['32', 'Produk', 'Menghapus produk', 'Produk terhapus dari daftar', 'Produk berhasil dihapus', 'Valid'],
        # === PEMBELIAN ===
        ['33', 'Pembelian', 'Menampilkan daftar pembelian dengan filter', 'Daftar tampil, filter supplier & tanggal berfungsi', 'Filter berfungsi menyaring data', 'Valid'],
        ['34', 'Pembelian', 'Membuat transaksi pembelian baru', 'Transaksi tersimpan, nomor PO auto-generated, stok bertambah', 'PO tersimpan, stok varian bertambah sesuai qty', 'Valid'],
        ['35', 'Pembelian', 'Membuat pembelian tanpa memilih supplier', 'Validasi error, pembelian tidak tersimpan', 'Pesan validasi "Supplier harus dipilih"', 'Valid'],
        ['36', 'Pembelian', 'Melihat detail pembelian', 'Menampilkan info supplier, item, harga, subtotal, total', 'Detail pembelian menampilkan data lengkap', 'Valid'],
        ['37', 'Pembelian', 'Notifikasi email ke supplier setelah pembelian', 'Email otomatis terkirim ke supplier', 'Email terkirim (tercatat di log)', 'Valid'],
        # === PENJUALAN ===
        ['38', 'Penjualan', 'Menampilkan daftar penjualan dengan filter', 'Daftar tampil, filter tanggal berfungsi', 'Filter berfungsi menyaring data', 'Valid'],
        ['39', 'Penjualan', 'Membuat transaksi penjualan baru', 'Transaksi tersimpan, invoice auto-generated, stok berkurang', 'Invoice tersimpan, stok varian berkurang', 'Valid'],
        ['40', 'Penjualan', 'Scan barcode pada form penjualan', 'Produk otomatis ditambahkan ke keranjang', 'Produk dengan barcode sesuai langsung masuk keranjang', 'Valid'],
        ['41', 'Penjualan', 'Scan barcode yang tidak terdaftar', 'Menampilkan pesan "Produk tidak ditemukan"', 'Menampilkan pesan error', 'Valid'],
        ['42', 'Penjualan', 'Penjualan dengan qty melebihi stok', 'Validasi error, penjualan tidak tersimpan', 'Pesan "Stok tidak mencukupi"', 'Valid'],
        ['43', 'Penjualan', 'Melihat detail penjualan', 'Menampilkan info kasir, item terjual, harga, total', 'Detail penjualan menampilkan data lengkap', 'Valid'],
        ['44', 'Penjualan', 'Notifikasi reorder alert saat stok ≤ reorder level', 'Email otomatis terkirim ke admin', 'Email reorder alert terkirim (tercatat di log)', 'Valid'],
        # === RETUR ===
        ['45', 'Retur', 'Menampilkan daftar retur dengan filter', 'Daftar tampil, filter supplier & tanggal berfungsi', 'Filter berfungsi menyaring data', 'Valid'],
        ['46', 'Retur', 'Membuat retur baru: pilih PO, item, qty & alasan', 'Retur tersimpan, stok berkurang', 'Retur tersimpan, stok varian berkurang', 'Valid'],
        ['47', 'Retur', 'Membuat retur tanpa alasan', 'Retur tetap tersimpan (alasan opsional)', 'Retur tersimpan tanpa alasan', 'Valid'],
        ['48', 'Retur', 'Melihat detail retur', 'Menampilkan info supplier, ref PO, item, alasan', 'Detail retur menampilkan data lengkap', 'Valid'],
        # === STOK OPNAME ===
        ['49', 'Stok Opname', 'Menampilkan daftar stok opname', 'Daftar tampil dengan status draft/confirmed', 'Menampilkan 2 opname', 'Valid'],
        ['50', 'Stok Opname', 'Memulai opname baru: input stok fisik', 'Data tersimpan otomatis sebagai draft, selisih dihitung', 'Stok fisik tersimpan, selisih tampil real-time', 'Valid'],
        ['51', 'Stok Opname', 'Scan barcode pada form opname', 'Baris varian yang sesuai di-highlight', 'Varian yang di-scan langsung disorot', 'Valid'],
        ['52', 'Stok Opname', 'Konfirmasi opname (draft → confirmed)', 'Status confirmed, stok varian disesuaikan', 'Stok varian diperbarui sesuai stok fisik', 'Valid'],
        ['53', 'Stok Opname', 'Melihat detail opname', 'Menampilkan ringkasan dan tabel selisih', 'Detail menampilkan total varian, selisih', 'Valid'],
        # === LAPORAN ===
        ['54', 'Laporan', 'Mengakses indeks laporan', 'Menampilkan 5 kartu laporan sesuai role', '5 kartu laporan tampil untuk admin', 'Valid'],
        ['55', 'Laporan', 'Laporan Stok: filter kategori & pencarian', 'Data stok tampil sesuai filter, ringkasan akurat', 'Filter berfungsi, ringkasan sesuai', 'Valid'],
        ['56', 'Laporan', 'Laporan Penjualan: filter rentang tanggal', 'Data penjualan sesuai rentang, total akurat', 'Filter tanggal berfungsi', 'Valid'],
        ['57', 'Laporan', 'Laporan Pembelian: filter supplier & tanggal', 'Data pembelian tampil sesuai filter', 'Filter berfungsi', 'Valid'],
        ['58', 'Laporan', 'Laporan Retur: filter supplier & tanggal', 'Data retur tampil sesuai filter', 'Filter berfungsi', 'Valid'],
        ['59', 'Laporan', 'Laporan Stok Opname: filter status & tanggal', 'Data opname tampil sesuai filter', 'Filter berfungsi, detail selisih per varian tampil', 'Valid'],
        ['60', 'Laporan', 'Ekspor laporan ke Excel', 'File .xlsx terdownload dengan data sesuai', 'File Excel terdownload, data sesuai laporan', 'Valid'],
        # === AKUN ===
        ['61', 'Akun', 'Menampilkan daftar akun', 'Semua pengguna tampil dengan data lengkap', 'Menampilkan akun terdaftar', 'Valid'],
        ['62', 'Akun', 'Menambah akun baru dengan role kasir', 'Akun tersimpan, langsung aktif, muncul di daftar', 'Akun baru muncul dengan status Aktif', 'Valid'],
        ['63', 'Akun', 'Menambah akun dengan email sudah terdaftar', 'Validasi error, email harus unik', 'Pesan "Email sudah terdaftar"', 'Valid'],
        ['64', 'Akun', 'Mengedit akun: ubah nama dan role', 'Data akun berubah dan tersimpan', 'Nama dan role berhasil diubah', 'Valid'],
        ['65', 'Akun', 'Menonaktifkan akun (toggle active)', 'Status Nonaktif, user tidak bisa login', 'Akun nonaktif, login ditolak', 'Valid'],
        ['66', 'Akun', 'Mengaktifkan kembali akun nonaktif', 'Status Aktif, user bisa login kembali', 'Akun aktif, login berhasil', 'Valid'],
        # === PROFIL ===
        ['67', 'Profil', 'Mengakses halaman profil', 'Menampilkan form info profil & ganti password', 'Kedua form tampil', 'Valid'],
        ['68', 'Profil', 'Memperbarui nama profil', 'Nama berubah dan tersimpan', 'Nama berhasil diperbarui', 'Valid'],
        ['69', 'Profil', 'Mengganti password dengan konfirmasi benar', 'Password berubah, bisa login dengan password baru', 'Login dengan password baru berhasil', 'Valid'],
        ['70', 'Profil', 'Mengganti password dengan konfirmasi tidak cocok', 'Validasi error, password tidak berubah', 'Pesan "Konfirmasi password tidak cocok"', 'Valid'],
        # === SCANNER MOBILE ===
        ['71', 'Scanner Mobile', 'Membuat token scanner dari halaman penjualan', 'QR code tampil dengan token unik', 'QR code berisi token tampil di modal', 'Valid'],
        ['72', 'Scanner Mobile', 'Membuka URL scanner di perangkat mobile', 'Halaman scanner dengan antarmuka kamera', 'Scanner tampil dengan badge "Menunggu scan..."', 'Valid'],
        ['73', 'Scanner Mobile', 'Memindai barcode produk terdaftar', 'Data barcode terkirim ke desktop', 'Barcode terkirim, toast "Produk ditemukan" tampil', 'Valid'],
        ['74', 'Scanner Mobile', 'Memindai barcode tidak terdaftar', 'Pesan "Produk tidak ditemukan" di mobile', 'Status error "Produk tidak ditemukan"', 'Valid'],
        ['75', 'Scanner Mobile', 'Koneksi terputus (desktop logout/refresh)', 'Overlay "Sesi Habis" tampil di mobile', 'Overlay tampil, scanner berhenti', 'Valid'],
    ]
)
add_figure_caption(doc, 'Tabel 4.1 Hasil Pengujian Black Box Sistem Informasi Kasyraa.co')

# --- Kesimpulan Pengujian ---
add_para(doc, 'Berdasarkan hasil pengujian black box yang telah dilakukan terhadap seluruh modul sistem, diperoleh kesimpulan sebagai berikut:')
add_para(doc, '1. Seluruh skenario pengujian pada 14 modul menghasilkan status "Valid", yang menunjukkan bahwa fungsionalitas sistem berjalan sesuai dengan kebutuhan yang telah dirancang.')
add_para(doc, '2. Total skenario pengujian yang dilakukan sebanyak 75 skenario, meliputi modul login (4), dashboard (4), kategori (5), satuan (4), supplier (5), produk (10), pembelian (5), penjualan (7), retur (4), stok opname (5), laporan (7), manajemen akun (6), profil (4), dan scanner mobile (5).')
add_para(doc, '3. Sistem mampu menangani input valid maupun invalid dengan memberikan respons yang sesuai, seperti menampilkan pesan validasi, pesan kesalahan, atau mengarahkan pengguna ke halaman yang tepat.')
add_para(doc, '4. Seluruh fitur CRUD pada modul data master berfungsi dengan baik, dan transaksi pembelian/penjualan/retur/opname dapat memperbarui stok secara akurat.')

# --- Penutup Bab ---
add_para(doc, '')
add_para(doc, 'Berdasarkan hasil implementasi yang telah dipaparkan, seluruh modul dalam sistem informasi inventori dan point of sale Kasyraa.co telah berhasil dikembangkan dan berjalan sesuai dengan kebutuhan fungsional yang telah dirancang pada Bab 3. Sistem menyediakan antarmuka yang intuitif dan responsif untuk mendukung operasional bisnis Kasyraa.co dalam mengelola inventori, transaksi pembelian, penjualan, retur, stok opname, serta pelaporan.')

# Save
doc.save(OUTPUT_PATH)
print(f"DOCX saved to: {OUTPUT_PATH}")
print("Done!")
