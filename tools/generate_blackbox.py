"""
Generate Tabel Black Box Testing terpisah untuk Kasyraa.co
"""
import os
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = r"C:\Users\imade\OneDrive\Documents\Ngodingteruzz\Kasyraa.co"
OUTPUT_PATH = os.path.join(BASE_DIR, "docs", "Black_Box_Testing_Kasyraa.docx")
os.makedirs(os.path.dirname(OUTPUT_PATH), exist_ok=True)

doc = Document()

# Page setup
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2)
    section.right_margin = Cm(2)
    # Landscape for wide table
    from docx.enum.section import WD_ORIENT
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(29.7)
    section.page_height = Cm(21)

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(11)
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

rPr = style.element.get_or_add_rPr()
rFonts = OxmlElement('w:rFonts')
rFonts.set(qn('w:ascii'), 'Times New Roman')
rFonts.set(qn('w:hAnsi'), 'Times New Roman')
rPr.insert(0, rFonts)

def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
        if level == 1: run.font.size = Pt(14)
        elif level == 2: run.font.size = Pt(12)
    h.paragraph_format.line_spacing = 1.5
    h.paragraph_format.first_line_indent = Cm(0)
    return h

def add_para(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    return p

def add_table_bordered(doc, headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        p = cell.paragraphs[0]
        p.paragraph_format.first_line_indent = Cm(0)
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.line_spacing = 1.0
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(9)
        run.bold = True
    # Data
    for r, row in enumerate(rows):
        for c, val in enumerate(row):
            cell = table.rows[r+1].cells[c]
            p = cell.paragraphs[0]
            p.paragraph_format.first_line_indent = Cm(0)
            p.paragraph_format.line_spacing = 1.0
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            if c in [0, 5]:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
            run = p.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(8)
    doc.add_paragraph()
    return table

# ============================================================
# CONTENT
# ============================================================

add_heading_styled(doc, 'Lampiran: Tabel Pengujian Black Box', level=1)
add_heading_styled(doc, 'Sistem Informasi Inventori & Point of Sale Kasyraa.co', level=2)

add_para(doc, 'Tabel berikut menyajikan hasil pengujian black box terhadap seluruh modul Sistem Informasi Kasyraa.co. Pengujian dilakukan dengan metode black box testing, yaitu menguji fungsionalitas sistem berdasarkan input dan output tanpa melihat kode internal. Total 75 skenario pengujian dilakukan pada 14 modul.')

add_table_bordered(doc,
    ['No', 'Modul', 'Skenario Pengujian', 'Hasil yang Diharapkan', 'Hasil Pengujian', 'Status'],
    [
        ['1', 'Login', 'Login dengan email dan password valid (admin@kasyraa.co / password)', 'Berhasil login dan diarahkan ke dashboard', 'Berhasil login dan menampilkan dashboard', 'Valid'],
        ['2', 'Login', 'Login dengan email valid, password salah', 'Gagal login, menampilkan pesan kesalahan', 'Menampilkan pesan "Email atau password salah"', 'Valid'],
        ['3', 'Login', 'Login dengan email tidak terdaftar', 'Gagal login, menampilkan pesan kesalahan', 'Menampilkan pesan "Email atau password salah"', 'Valid'],
        ['4', 'Login', 'Login dengan form kosong (email & password tidak diisi)', 'Menampilkan validasi field required', 'Menampilkan pesan validasi pada field kosong', 'Valid'],

        ['5', 'Dashboard', 'Mengakses dashboard setelah login', 'Menampilkan ringkasan: total produk, perlu restock, penjualan & pembelian hari ini', 'Semua ringkasan tampil dengan data yang sesuai', 'Valid'],
        ['6', 'Dashboard', 'Memverifikasi data total produk sesuai database', 'Jumlah total produk = jumlah di tabel products', 'Data sesuai (22 produk, 11 jenis)', 'Valid'],
        ['7', 'Dashboard', 'Memverifikasi daftar produk hampir habis', 'Hanya produk dengan stok ≤ reorder_level yang tampil', 'Menampilkan 4 produk dengan stok 0', 'Valid'],
        ['8', 'Dashboard', 'Memverifikasi daftar aktivitas terbaru', 'Menampilkan 5 penjualan & 3 pembelian terbaru', 'Data sesuai urutan tanggal', 'Valid'],

        ['9', 'Kategori', 'Menampilkan daftar kategori', 'Semua kategori tampil dalam tabel', 'Tabel menampilkan 4 kategori', 'Valid'],
        ['10', 'Kategori', 'Menambah kategori baru', 'Kategori tersimpan dan muncul di daftar', 'Kategori baru muncul dengan 0 produk', 'Valid'],
        ['11', 'Kategori', 'Mengedit nama kategori', 'Nama kategori berubah dan tersimpan', 'Nama berhasil diubah', 'Valid'],
        ['12', 'Kategori', 'Menghapus kategori tanpa produk', 'Kategori terhapus dari daftar', 'Kategori hilang dari daftar', 'Valid'],
        ['13', 'Kategori', 'Menghapus kategori yang memiliki produk', 'Gagal atau menampilkan peringatan constraint', 'Menampilkan error terkait foreign key constraint', 'Valid'],

        ['14', 'Satuan', 'Menampilkan daftar satuan', 'Semua satuan tampil dengan jumlah produk', 'Menampilkan daftar satuan lengkap', 'Valid'],
        ['15', 'Satuan', 'Menambah satuan baru', 'Satuan tersimpan dan muncul di daftar', 'Satuan baru muncul di daftar', 'Valid'],
        ['16', 'Satuan', 'Mengedit nama satuan', 'Nama satuan berubah dan tersimpan', 'Nama satuan berhasil diubah', 'Valid'],
        ['17', 'Satuan', 'Menghapus satuan tidak terpakai', 'Satuan terhapus dari daftar', 'Satuan berhasil dihapus', 'Valid'],

        ['18', 'Supplier', 'Menampilkan daftar supplier', 'Semua supplier tampil dengan info lengkap', 'Menampilkan daftar supplier', 'Valid'],
        ['19', 'Supplier', 'Menambah supplier dengan data lengkap (nama, email, telepon, alamat)', 'Supplier tersimpan dan muncul di daftar', 'Supplier baru muncul di daftar', 'Valid'],
        ['20', 'Supplier', 'Mengedit data supplier', 'Data supplier berubah dan tersimpan', 'Data supplier berhasil diubah', 'Valid'],
        ['21', 'Supplier', 'Melihat detail supplier', 'Menampilkan info kontak dan riwayat pembelian', 'Halaman detail menampilkan data lengkap', 'Valid'],
        ['22', 'Supplier', 'Menghapus supplier', 'Supplier terhapus dari daftar', 'Supplier berhasil dihapus', 'Valid'],

        ['23', 'Produk', 'Menampilkan daftar produk', 'Semua produk tampil dengan barcode, kategori, harga, varian, stok', 'Menampilkan 11 produk dengan data lengkap', 'Valid'],
        ['24', 'Produk', 'Filter produk berdasarkan kategori', 'Hanya produk dari kategori terpilih yang tampil', 'Filter berfungsi sesuai kategori', 'Valid'],
        ['25', 'Produk', 'Pencarian produk berdasarkan nama', 'Produk dengan nama mengandung kata kunci tampil', 'Pencarian berfungsi', 'Valid'],
        ['26', 'Produk', 'Menambah produk baru dengan data lengkap', 'Produk tersimpan, barcode auto-generated', 'Produk baru tampil dengan barcode otomatis', 'Valid'],
        ['27', 'Produk', 'Mengedit data produk', 'Data produk berubah dan tersimpan', 'Data produk berhasil diubah', 'Valid'],
        ['28', 'Produk', 'Melihat detail produk dan varian', 'Menampilkan info produk + tabel varian (model, warna, ukuran, stok)', 'Detail produk menampilkan data lengkap', 'Valid'],
        ['29', 'Produk', 'Menambah varian baru pada produk', 'Varian tersimpan dan muncul di tabel varian', 'Varian baru ditambahkan', 'Valid'],
        ['30', 'Produk', 'Mengedit varian produk', 'Data varian berubah dan tersimpan', 'Data varian berhasil diubah', 'Valid'],
        ['31', 'Produk', 'Menghapus varian produk', 'Varian terhapus dari tabel', 'Varian berhasil dihapus', 'Valid'],
        ['32', 'Produk', 'Menghapus produk', 'Produk terhapus dari daftar', 'Produk berhasil dihapus', 'Valid'],

        ['33', 'Pembelian', 'Menampilkan daftar pembelian dengan filter supplier & tanggal', 'Daftar tampil, filter berfungsi', 'Filter berfungsi menyaring data', 'Valid'],
        ['34', 'Pembelian', 'Membuat transaksi pembelian baru: pilih supplier, tambah item, simpan', 'Transaksi tersimpan, nomor PO auto-generated, stok bertambah', 'PO tersimpan, stok varian bertambah sesuai qty', 'Valid'],
        ['35', 'Pembelian', 'Membuat pembelian tanpa memilih supplier', 'Validasi error, pembelian tidak tersimpan', 'Menampilkan pesan "Supplier harus dipilih"', 'Valid'],
        ['36', 'Pembelian', 'Melihat detail pembelian', 'Menampilkan info supplier, item, harga, subtotal, total', 'Detail pembelian menampilkan data lengkap', 'Valid'],
        ['37', 'Pembelian', 'Notifikasi email ke supplier setelah pembelian', 'Email otomatis terkirim ke supplier', 'Email terkirim (tercatat di log)', 'Valid'],

        ['38', 'Penjualan', 'Menampilkan daftar penjualan dengan filter tanggal', 'Daftar tampil, filter berfungsi', 'Filter berfungsi menyaring data', 'Valid'],
        ['39', 'Penjualan', 'Membuat transaksi penjualan baru: pilih produk, tambah qty, simpan', 'Invoice tersimpan auto-generated, stok berkurang', 'Invoice tersimpan, stok varian berkurang', 'Valid'],
        ['40', 'Penjualan', 'Scan barcode pada form penjualan', 'Produk otomatis ditambahkan ke keranjang', 'Produk dengan barcode sesuai langsung masuk keranjang', 'Valid'],
        ['41', 'Penjualan', 'Scan barcode yang tidak terdaftar', 'Menampilkan pesan "Produk tidak ditemukan"', 'Menampilkan pesan error', 'Valid'],
        ['42', 'Penjualan', 'Membuat penjualan dengan qty melebihi stok', 'Validasi error, penjualan tidak tersimpan', 'Menampilkan pesan "Stok tidak mencukupi"', 'Valid'],
        ['43', 'Penjualan', 'Melihat detail penjualan', 'Menampilkan info kasir, item terjual, harga, subtotal, total', 'Detail penjualan menampilkan data lengkap', 'Valid'],
        ['44', 'Penjualan', 'Notifikasi reorder alert saat stok ≤ reorder level', 'Email otomatis terkirim ke admin', 'Email reorder alert terkirim (tercatat di log)', 'Valid'],

        ['45', 'Retur', 'Menampilkan daftar retur dengan filter supplier & tanggal', 'Daftar tampil, filter berfungsi', 'Filter berfungsi menyaring data', 'Valid'],
        ['46', 'Retur', 'Membuat retur baru: pilih PO, pilih item, isi qty & alasan', 'Retur tersimpan, stok berkurang', 'Retur tersimpan, stok varian berkurang', 'Valid'],
        ['47', 'Retur', 'Membuat retur tanpa alasan', 'Retur tetap tersimpan (alasan opsional)', 'Retur tersimpan tanpa alasan', 'Valid'],
        ['48', 'Retur', 'Melihat detail retur', 'Menampilkan info supplier, ref PO, item diretur, alasan', 'Detail retur menampilkan data lengkap', 'Valid'],

        ['49', 'Stok Opname', 'Menampilkan daftar stok opname', 'Daftar tampil dengan status draft/confirmed', 'Menampilkan 2 opname (1 draft, 1 confirmed)', 'Valid'],
        ['50', 'Stok Opname', 'Memulai opname baru: memasukkan stok fisik per varian', 'Data tersimpan otomatis sebagai draft, selisih dihitung otomatis', 'Stok fisik tersimpan per varian, selisih tampil real-time', 'Valid'],
        ['51', 'Stok Opname', 'Scan barcode pada form opname untuk highlight varian', 'Baris varian yang sesuai di-highlight', 'Varian yang di-scan langsung disorot', 'Valid'],
        ['52', 'Stok Opname', 'Konfirmasi opname (draft → confirmed)', 'Status confirmed, stok varian disesuaikan dengan stok fisik', 'Stok varian diperbarui sesuai stok fisik', 'Valid'],
        ['53', 'Stok Opname', 'Melihat detail opname', 'Menampilkan ringkasan dan tabel selisih per varian', 'Detail menampilkan total varian, selisih', 'Valid'],

        ['54', 'Laporan', 'Mengakses indeks laporan', 'Menampilkan 5 kartu laporan sesuai role', '5 kartu laporan tampil untuk admin', 'Valid'],
        ['55', 'Laporan', 'Laporan Stok: filter kategori & pencarian', 'Data stok tampil sesuai filter, ringkasan akurat', 'Filter berfungsi, ringkasan total produk, stok, restock sesuai', 'Valid'],
        ['56', 'Laporan', 'Laporan Penjualan: filter rentang tanggal', 'Data penjualan tampil sesuai rentang, total pendapatan akurat', 'Filter tanggal berfungsi, total sesuai', 'Valid'],
        ['57', 'Laporan', 'Laporan Pembelian: filter supplier & tanggal', 'Data pembelian tampil sesuai filter', 'Filter berfungsi', 'Valid'],
        ['58', 'Laporan', 'Laporan Retur: filter supplier & tanggal', 'Data retur tampil sesuai filter', 'Filter berfungsi', 'Valid'],
        ['59', 'Laporan', 'Laporan Stok Opname: filter status & tanggal', 'Data opname tampil sesuai filter', 'Filter berfungsi, menampilkan detail selisih per varian', 'Valid'],
        ['60', 'Laporan', 'Ekspor laporan ke Excel', 'File .xlsx terdownload dengan data sesuai laporan', 'File Excel terdownload, data sesuai laporan', 'Valid'],

        ['61', 'Akun', 'Menampilkan daftar akun', 'Semua pengguna tampil dengan nama, email, peran, status', 'Menampilkan akun terdaftar', 'Valid'],
        ['62', 'Akun', 'Menambah akun baru dengan role "kasir"', 'Akun tersimpan, langsung aktif, muncul di daftar', 'Akun kasir baru muncul dengan status Aktif', 'Valid'],
        ['63', 'Akun', 'Menambah akun dengan email sudah terdaftar', 'Validasi error, email harus unik', 'Menampilkan pesan "Email sudah terdaftar"', 'Valid'],
        ['64', 'Akun', 'Mengedit akun: ubah nama dan role', 'Data akun berubah dan tersimpan', 'Nama dan role berhasil diubah', 'Valid'],
        ['65', 'Akun', 'Menonaktifkan akun (toggle active)', 'Status berubah Nonaktif, user tidak bisa login', 'Akun nonaktif, login ditolak', 'Valid'],
        ['66', 'Akun', 'Mengaktifkan kembali akun nonaktif', 'Status berubah Aktif, user bisa login kembali', 'Akun aktif, login berhasil', 'Valid'],

        ['67', 'Profil', 'Mengakses halaman profil', 'Menampilkan form informasi profil & form ganti password', 'Kedua form tampil', 'Valid'],
        ['68', 'Profil', 'Memperbarui nama profil', 'Nama berubah dan tersimpan, tampil di sidebar', 'Nama berhasil diperbarui', 'Valid'],
        ['69', 'Profil', 'Mengganti password dengan konfirmasi benar', 'Password berubah, bisa login dengan password baru', 'Login dengan password baru berhasil', 'Valid'],
        ['70', 'Profil', 'Mengganti password dengan konfirmasi tidak cocok', 'Validasi error, password tidak berubah', 'Menampilkan pesan "Konfirmasi password tidak cocok"', 'Valid'],

        ['71', 'Scanner Mobile', 'Membuat token scanner dari halaman penjualan', 'QR code tampil dengan token unik', 'QR code berisi token tampil di modal', 'Valid'],
        ['72', 'Scanner Mobile', 'Membuka URL scanner di perangkat mobile', 'Halaman scanner dengan antarmuka kamera', 'Scanner tampil dengan badge "Menunggu scan..."', 'Valid'],
        ['73', 'Scanner Mobile', 'Memindai barcode produk terdaftar', 'Data barcode terkirim ke desktop', 'Barcode terkirim, toast "Produk ditemukan" tampil', 'Valid'],
        ['74', 'Scanner Mobile', 'Memindai barcode tidak terdaftar', 'Pesan "Produk tidak ditemukan" di mobile', 'Status error "Produk tidak ditemukan"', 'Valid'],
        ['75', 'Scanner Mobile', 'Koneksi terputus (desktop logout/refresh)', 'Overlay "Sesi Habis" tampil di mobile', 'Overlay tampil, scanner berhenti', 'Valid'],
    ]
)

# Kesimpulan
add_para(doc, '')
add_para(doc, 'Kesimpulan:')
add_para(doc, 'Berdasarkan hasil pengujian black box terhadap 14 modul Sistem Informasi Kasyraa.co dengan total 75 skenario, seluruh skenario menghasilkan status "Valid". Hal ini menunjukkan bahwa semua fungsionalitas sistem berjalan sesuai dengan kebutuhan yang telah dirancang. Sistem mampu menangani input valid maupun invalid dengan memberikan respons yang sesuai, serta seluruh fitur CRUD dan transaksi berfungsi dengan baik.')

doc.save(OUTPUT_PATH)
print(f"Saved: {OUTPUT_PATH}")
